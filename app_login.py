# -*- coding: utf-8 -*-
import json
import os
import queue
import re
import shutil
import sys
import tempfile
import threading
import time
from datetime import datetime, timedelta
from tkinter import (
    END, Button, Frame, Label, Tk, filedialog, font as tkfont, messagebox, ttk
)

import requests
import undetected_chromedriver as uc
from PIL import Image, ImageTk
from selenium.common.exceptions import TimeoutException, WebDriverException
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

try:
    import docx
    from utils.obtertoken import obter_token
except ImportError:
    messagebox.showerror(
        "Bibliotecas Faltando",
        "A biblioteca 'python-docx' é necessária.\n"
        "Por favor, instale-a com o comando:\n\n"
        "pip install python-docx"
    )
    sys.exit()

def resource_path(relative_path):
    """ Obtém o caminho absoluto para o recurso, funciona para dev e para PyInstaller """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# Fila para comunicação entre a automação e a interface
RESULTADO_QUEUE = queue.Queue()


class AppPrincipal(Tk):
    """Classe principal da aplicação de análise de processos."""

    def __init__(self):
        super().__init__()
        self.title("Sistema de Análise de Processos")
        self.geometry("950x700")
        self.resizable(True, True)
        self.main_driver = None
        self.profile_path = None
        self.access_token = None
        self.dados_processos = None
        self.resultados_finais = []
        self.treeview_resultados = None
        self.status_label = None
        self._setup_ui()
        self.protocol("WM_DELETE_WINDOW", self.ao_fechar_janela)

    def _setup_ui(self):
        """Configura a interface gráfica inicial."""
        logo_frame = Frame(self)
        logo_frame.pack(pady=10)
        try:
            caminho_imagem = resource_path("logo_detran.png")
            img = Image.open(caminho_imagem)
            img.thumbnail((200, 100))
            self.logo_image = ImageTk.PhotoImage(img)
            logo_label = Label(logo_frame, image=self.logo_image)
            logo_label.pack()
        except Exception as e:
            Label(logo_frame, text="[Imagem 'logo_detran.png' não encontrada]", fg="grey").pack()
            print(f"ERRO AO CARREGAR IMAGEM: {e}")
        separator = Frame(self, height=2, bd=1, relief="sunken")
        separator.pack(fill="x", padx=20, pady=5)
        self.container = Frame(self)
        self.container.pack(expand=True, fill="both", padx=30, pady=20)
        self.mostrar_tela_login()

    def limpar_tela(self):
        """Remove todos os widgets do container principal."""
        for widget in self.container.winfo_children():
            widget.destroy()

    def mostrar_tela_login(self):
        """Exibe a tela inicial de login."""
        self.limpar_tela()
        self.title("Sistema - Autenticação")
        self.resultados_finais.clear()
        self.treeview_resultados = None
        bold_font = tkfont.Font(family="Helvetica", size=14, weight="bold")
        Label(self.container, text="Acesse o sistema para continuar", font=bold_font).pack(pady=20)
        Button(
            self.container, text="Fazer Login",
            font=tkfont.Font(family="Helvetica", size=16, weight="bold"),
            command=self.iniciar_processo_login_token,
            bg="#1976D2", fg="white", pady=20, width=20
        ).pack(pady=20)

    def iniciar_processo_login_token(self):
        """Inicia a thread para o processo de login."""
        self.mostrar_tela_aguardando_login()
        threading.Thread(target=self._worker_login_token, daemon=True).start()
        self.verificar_resultado_login_token()

    def mostrar_tela_aguardando_login(self):
        """Exibe tela de espera durante o login."""
        self.limpar_tela()
        bold_font = tkfont.Font(family="Helvetica", size=12, weight="bold")
        texto_aguarde = (
            "Iniciando navegador...\n\n"
            "A primeira execução pode demorar um pouco mais.\n"
            "Por favor, aguarde a janela do Chrome abrir."
        )
        Label(self.container, text=texto_aguarde, font=bold_font).pack(pady=50)

    def _worker_login_token(self):
        """Executa a automação para obter o token de acesso."""
        try:
            if not self.main_driver:
                self.profile_path = os.path.join(tempfile.gettempdir(), f"chrome_profile_{int(time.time())}")
                options = uc.ChromeOptions()
                options.add_argument(f'--user-data-dir={self.profile_path}')
                self.main_driver = uc.Chrome(options=options, headless=False, use_subprocess=True)
                self.main_driver.maximize_window()
                self.main_driver.execute_script("window.focus();")
            client_id = "e1463040-77cb-4e8a-ae78-7d9b6ecd4b6c"
            scope = "openid profile api-sigades-documento api-sigades-processo api-sigades-consultar"
            redirect_uri = "https://detran.es.gov.br/"
            auth_url = f"https://acessocidadao.es.gov.br/is/connect/authorize?response_type=code%20id_token&client_id={client_id}&scope={scope}&redirect_uri={redirect_uri}&nonce=1000"
            self.main_driver.get(auth_url)
            wait = WebDriverWait(self.main_driver, 60)
            try:
                login_classico = wait.until(EC.element_to_be_clickable((By.XPATH, "//span[contains(text(), 'Login Clássico')]")))
                login_classico.click()
            except TimeoutException:
                print("ℹ️ Link 'Login Clássico' não encontrado. Prosseguindo.")
            print("✅ Página de login pronta. Por favor, faça o login no navegador.")
            wait.until(EC.url_contains("code="))
            print("✅ Login detectado com sucesso!")
            url_final = self.main_driver.current_url
            code_match = re.search(r'code=([^&#]+)', url_final)
            if code_match:
                auth_code = code_match.group(1)
                token_info = obter_token(code=auth_code, silent_mode=False, base="homologacao")
                if token_info and "access_token" in token_info:
                    RESULTADO_QUEUE.put({"status": "sucesso_login_token", "token_info": token_info})
                else:
                    raise Exception(f"Falha ao obter o token de acesso. Resposta: {token_info}")
            else:
                raise Exception("'code' não encontrado na URL final.")
        except Exception as e:
            print(f"ERRO no worker de login: {e}")
            RESULTADO_QUEUE.put({"status": "falha_login_token", "erro": str(e)})

    def verificar_resultado_login_token(self):
        """Verifica o resultado do processo de login na fila."""
        try:
            resultado = RESULTADO_QUEUE.get_nowait()
            if resultado["status"] == "sucesso_login_token":
                self.access_token = resultado["token_info"].get("access_token")
                print("✅ Token obtido com sucesso.")
                self.iniciar_busca_processos()
            else:
                messagebox.showerror("Falha no Login", f"{resultado.get('erro', 'Erro desconhecido.')}\nTente novamente.")
                self.mostrar_tela_login()
        except queue.Empty:
            self.after(200, self.verificar_resultado_login_token)

    def iniciar_busca_processos(self):
        self.mostrar_tela_buscando_processos()
        threading.Thread(target=self._worker_buscar_processos, daemon=True).start()
        self.verificar_resultado_api()

    def mostrar_tela_buscando_processos(self):
        self.limpar_tela()
        self.title("Sistema - Buscando Dados")
        bold_font = tkfont.Font(family="Helvetica", size=12, weight="bold")
        Label(self.container, text="Autenticação concluída!\nBuscando processos na API...", font=bold_font).pack(pady=50)

    def _worker_buscar_processos(self):
        try:
            if not self.access_token:
                raise Exception("Token de acesso é inválido.")
            url = "https://api.e-docs.es.gov.br/v2/processos/paginated-search"
            headers = {'accept': 'application/json', 'Authorization': f'Bearer {self.access_token}', 'Content-Type': 'application/json-patch+json'}
            payload = {"idsLocaisAtuais": ["b7926835-d3ac-48c9-a1a5-320c85b98a24"], "somenteEmAndamento": True, "pagina": 1, "tamanhoPagina": 100}
            response = requests.post(url, headers=headers, json=payload, timeout=30)
            response.raise_for_status()
            RESULTADO_QUEUE.put({"status": "sucesso_api", "dados": response.json()})
        except Exception as e:
            print(f"ERRO na busca de processos: {e}")
            RESULTADO_QUEUE.put({"status": "falha_api", "erro": str(e)})

    def verificar_resultado_api(self):
        try:
            resultado = RESULTADO_QUEUE.get_nowait()
            if resultado["status"] == "sucesso_api":
                self.dados_processos = resultado["dados"]
                processos = self.dados_processos.get('result', [])
                processos_filtrados = [p for p in processos if "DEFESA PRÉVIA MULTA" not in p.get('resumo', '')]
                self.dados_processos['result'] = processos_filtrados
                self.dados_processos['count'] = len(processos_filtrados)
                if self.dados_processos['count'] > 0:
                    print(f"✅ API OK. {self.dados_processos['count']} processos para automação.")
                    self.iniciar_login_renach()
                else:
                    messagebox.showinfo("Concluído", "Nenhum processo válido encontrado.")
                    self.mostrar_tela_login()
            else:
                messagebox.showerror("Falha na API", resultado.get('erro', 'Erro desconhecido.'))
                self.mostrar_tela_login()
        except queue.Empty:
            self.after(200, self.verificar_resultado_api)

    def iniciar_login_renach(self):
        self.mostrar_tela_login_renach()
        threading.Thread(target=self._worker_login_renach, daemon=True).start()
        self.verificar_resultado_login_renach()

    def mostrar_tela_login_renach(self):
        self.limpar_tela()
        self.title("Sistema - Login RENACH")
        bold_font = tkfont.Font(family="Helvetica", size=12, weight="bold")
        Label(self.container, text="Realizando login no portal RENACH...\n\nAguarde um momento.", font=bold_font).pack(pady=50)

    def _worker_login_renach(self):
        try:
            self.main_driver.maximize_window()
            self.main_driver.get("https://renach2.es.gov.br/")
            wait = WebDriverWait(self.main_driver, 20)
            wait.until(EC.element_to_be_clickable((By.XPATH, "//div[contains(text(), 'Entrar com o Acesso Cidadão')]"))).click()
            try:
                wait.until(EC.element_to_be_clickable((By.ID, "ctl00_conteudo_btnContinuarAcessoCidadao"))).click()
            except TimeoutException:
                pass
            time.sleep(2)
            RESULTADO_QUEUE.put({"status": "sucesso_login_renach"})
        except Exception as e:
            print(f"ERRO no login RENACH: {e}")
            RESULTADO_QUEUE.put({"status": "falha_login_renach", "erro": str(e)})

    def verificar_resultado_login_renach(self):
        try:
            resultado = RESULTADO_QUEUE.get_nowait()
            if resultado["status"] == "sucesso_login_renach":
                print("✅ Login no RENACH concluído.")
                self.mostrar_tela_resultados_em_progresso()
                self.iniciar_consulta_processo()
            else:
                messagebox.showerror("Falha no RENACH", resultado.get('erro', 'Erro desconhecido.'))
                self.mostrar_tela_login()
        except queue.Empty:
            self.after(200, self.verificar_resultado_login_renach)

    def mostrar_tela_resultados_em_progresso(self):
        self.limpar_tela()
        self.title("Sistema - Análise em Andamento...")
        frame_topo = Frame(self.container)
        frame_topo.pack(fill="x", pady=5)
        Button(
            frame_topo, text="📋 Gerar Relatório (.docx)", command=self.gerar_documento_word,
            font=tkfont.Font(size=12, weight="bold"), bg="#2E7D32", fg="white"
        ).pack(side="left")
        self.status_label = Label(frame_topo, text="Iniciando...", font=tkfont.Font(size=11))
        self.status_label.pack(side="right", padx=10)
        frame_tabela = Frame(self.container)
        frame_tabela.pack(expand=True, fill="both", pady=10)
        colunas = ("protocolo", "nome", "remessa", "status", "dias")
        self.treeview_resultados = ttk.Treeview(frame_tabela, columns=colunas, show="headings")
        headings = {
            "protocolo": ("Protocolo", 120), "nome": ("Nome do Condutor", 280),
            "remessa": ("Remessa SIT Próxima", 150), "status": ("Status", 100),
            "dias": ("Dias", 120)
        }
        for col, (text, width) in headings.items():
            self.treeview_resultados.heading(col, text=text)
            self.treeview_resultados.column(col, width=width, anchor="center")
        self.treeview_resultados.column("nome", anchor="w")
        scrollbar = ttk.Scrollbar(frame_tabela, orient="vertical", command=self.treeview_resultados.yview)
        self.treeview_resultados.configure(yscroll=scrollbar.set)
        scrollbar.pack(side="right", fill="y")
        self.treeview_resultados.pack(expand=True, fill="both")

    def iniciar_consulta_processo(self):
        """Controla o loop de automação, um processo por vez."""
        if not self.dados_processos or self.dados_processos['count'] == 0:
            # --- CORREÇÃO FINAL: FECHA O NAVEGADOR ANTES DO POP-UP ---
            if self.main_driver:
                print("✅ Automação concluída. Fechando o navegador...")
                try:
                    self.main_driver.quit()
                    self.main_driver = None
                except Exception as e:
                    print(f"ℹ️ Erro menor ao fechar o navegador: {e}")
            
            messagebox.showinfo("Fim da Fila", "Todos os processos foram concluídos!")
            total = len(self.resultados_finais)
            self.status_label.config(text=f"Concluído! {total} processos analisados.")
            return

        total_inicial = len(self.dados_processos.get('result', [])) + len(self.resultados_finais)
        self.status_label.config(text=f"Processando {len(self.resultados_finais) + 1} de {total_inicial}...")
        threading.Thread(target=self._worker_consultar_processo, daemon=True).start()
        self.verificar_resultado_consulta()

    def _worker_consultar_processo(self):
        try:
            protocolo = self.dados_processos['result'][0].get('protocolo')
            if not protocolo:
                raise Exception("Processo sem número de protocolo.")
            url_consulta = "https://renach2.es.gov.br/penalidade/mgp_consulta_processo_adm.aspx"
            self.main_driver.get(url_consulta)
            print(f"\n--- Consultando protocolo: {protocolo} ---")
            wait = WebDriverWait(self.main_driver, 20)
            wait.until(EC.element_to_be_clickable((By.ID, "ctl00_conteudo_ucPesquisa_txbNumProtocolo"))).send_keys(protocolo)
            wait.until(EC.element_to_be_clickable((By.ID, "ctl00_conteudo_ucPesquisa_btnFiltrar"))).click()
            nome_condutor = wait.until(EC.presence_of_element_located((By.ID, "ctl00_conteudo_ucDados_lblNomeTexto"))).text.strip()
            infracoes = self._extrair_infracoes(wait)
            dados_analisados = self._analisar_datas_infracoes(protocolo, nome_condutor, infracoes)
            RESULTADO_QUEUE.put({"status": "sucesso_consulta", "dados": dados_analisados})
        except Exception as e:
            print(f"ERRO na consulta do protocolo: {e}")
            RESULTADO_QUEUE.put({"status": "falha_consulta", "erro": str(e)})

    def _extrair_infracoes(self, wait):
        lista_infracoes = []
        try:
            tabela_id = "ctl00_conteudo_ucInfracoes_gvInfracoes"
            tabela = wait.until(EC.presence_of_element_located((By.ID, tabela_id)))
            linhas = tabela.find_elements(By.XPATH, ".//tr[contains(@class, 'RowStyle')]")
            for linha in linhas:
                celulas = linha.find_elements(By.TAG_NAME, "td")
                if len(celulas) >= 7:
                    lista_infracoes.append({"Remessa SIT": celulas[6].text.strip()})
        except TimeoutException:
            print("ℹ️ Nenhuma tabela de infrações encontrada.")
        return lista_infracoes
    
    def _analisar_datas_infracoes(self, protocolo, nome_condutor, infracoes):
        hoje = datetime.now()
        remessa_mais_proxima, menor_diff = None, float('inf')
        for infracao in infracoes:
            try:
                data_remessa = datetime.strptime(infracao["Remessa SIT"], '%d/%m/%Y')
                diff = abs((hoje - data_remessa).days)
                if diff < menor_diff:
                    menor_diff = diff
                    remessa_mais_proxima = data_remessa
            except (ValueError, KeyError):
                continue
        if remessa_mais_proxima:
            vencimento = remessa_mais_proxima + timedelta(days=360)
            remessa_str = remessa_mais_proxima.strftime('%d/%m/%Y')
            if vencimento.date() >= hoje.date():
                status = "Válida"
                dias = f"Faltam {(vencimento - hoje).days} dias"
            else:
                status = "Vencida"
                dias = f"Há {(hoje - vencimento).days} dias"
        else:
            remessa_str = "N/A"
            status = "Sem infrações" if not infracoes else "Erro de data"
            dias = ""
        return {"protocolo": protocolo, "nome": nome_condutor, "remessa_sit": remessa_str, "status": status, "dias": dias}

    def verificar_resultado_consulta(self):
        try:
            resultado = RESULTADO_QUEUE.get_nowait()
            if resultado["status"] == "sucesso_consulta":
                dados = resultado["dados"]
                self.resultados_finais.append(dados)
                if self.treeview_resultados:
                    self.treeview_resultados.insert("", END, values=list(dados.values()))
                self.dados_processos['result'].pop(0)
                self.dados_processos['count'] -= 1
                self.iniciar_consulta_processo()
            else:
                messagebox.showerror("Falha na Consulta", resultado.get('erro', 'Erro desconhecido.'))
                self.mostrar_tela_login()
        except queue.Empty:
            self.after(200, self.verificar_resultado_consulta)

    def gerar_documento_word(self):
        if not self.resultados_finais:
            messagebox.showwarning("Atenção", "Nenhum dado foi processado para gerar o relatório.")
            return
        try:
            filepath = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Documento Word", "*.docx")],
                title="Salvar Relatório Como...", initialfile="Relatorio_Processos.docx"
            )
            if not filepath: return
            doc = docx.Document()
            doc.add_heading('Relatório de Análise de Processos', 0)
            doc.add_paragraph(f"Relatório gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
            tabela = doc.add_table(rows=1, cols=5)
            tabela.style = 'Table Grid'
            hdr_cells = tabela.rows[0].cells
            col_names = ['Protocolo', 'Nome do Condutor', 'Remessa SIT', 'Status', 'Dias']
            for i, name in enumerate(col_names):
                hdr_cells[i].text = name
            for res in self.resultados_finais:
                row_cells = tabela.add_row().cells
                row_cells[0].text = res.get('protocolo', '')
                row_cells[1].text = res.get('nome', '')
                row_cells[2].text = res.get('remessa_sit', '')
                row_cells[3].text = res.get('status', '')
                row_cells[4].text = res.get('dias', '')
            doc.save(filepath)
            messagebox.showinfo("Sucesso", f"Relatório salvo com sucesso em:\n{filepath}")
        except Exception as e:
            messagebox.showerror("Erro ao Gerar Documento", f"Não foi possível criar o arquivo Word.\n\nErro: {e}")

    def ao_fechar_janela(self):
        """Encerra o driver e a aplicação de forma segura."""
        print("Fechando a aplicação...")
        if self.main_driver:
            try:
                self.main_driver.quit()
            except WebDriverException:
                print("ℹ️ Navegador já estava fechado.")
        if self.profile_path and os.path.exists(self.profile_path):
            shutil.rmtree(self.profile_path, ignore_errors=True)
            print("✅ Perfil temporário limpo.")
        self.destroy()


if __name__ == "__main__":
    app = AppPrincipal()
    app.mainloop()